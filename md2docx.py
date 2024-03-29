import subprocess
from pathlib import Path
import logging
from typing import Literal
import re
from rich.logging import RichHandler
from rich.traceback import install
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.shape import WD_INLINE_SHAPE
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Cm, Mm

CUSTOM_TEXT = ["1.Azure", "2.AWS", "3.WPEngine", "4.FraudWatch", "5.Cisco", "6.Barracuda", "7.Websites", "8.Summary"]  # pylint: disable=redefined-outer-name


class DirectoryInitializer:
    TOP_MARGIN = Cm(1)
    BOTTOM_MARGIN = Cm(1)
    LEFT_MARGIN = Cm(1)
    RIGHT_MARGIN = Cm(1)

    def __init__(self) -> None:
        self.input_dir = Path("input/CleanedTemplate")
        self.output_dir = Path("final_output")
        self.reference_dir = self.input_dir / "Reference"
        self.reference_doc = self.reference_dir / "refdoc.docx"
        self.custom_text = ["1.Azure", "2.AWS", "3.WPEngine", "4.FraudWatch", "5.Cisco", "6.Barracuda", "7.Websites", "8.Summary"]
        self._create_directories()

    def _create_directories(self) -> None:
        self.input_dir.mkdir(exist_ok=True)
        self.reference_dir.mkdir(exist_ok=True)
        self.output_dir.mkdir(exist_ok=True)


class DocxProcessor:

    def __init__(self, input_dir: str, output_dir: str, reference_dir: str, reference_doc: str) -> None:
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.reference_dir = Path(reference_dir)
        self.reference_doc = reference_doc

    def post_process_docx(self, doc_path) -> None:
        try:
            doc = Document(str(doc_path))
            formatter = DocumentFormatter(doc)
            formatter.set_document_font()
            formatter.set_margins(DirectoryInitializer.TOP_MARGIN, DirectoryInitializer.BOTTOM_MARGIN, DirectoryInitializer.LEFT_MARGIN,
                                  DirectoryInitializer.RIGHT_MARGIN)
            formatter.modify_document_styles()

            table_styler = TableStyler(doc)
            table_styler.autofit_tables_to_window()

            style_applier = StyleApplier(doc, table_styler)
            style_applier.apply_custom_styles()

            # Image Resizing
            image_resizer = ImageResizer(doc)
            image_resizer.autofit_images_to_window()

            # Section Management
            section_manager = SectionManager(doc)
            section_manager.keep_sections_together()
            section_manager.add_page_break_before_section(CUSTOM_TEXT)

            doc.save(str(doc_path))
            logger.info('Document processing completed.')
        except Exception as e:
            logger.error(f"Error during post-processing: {e}", exc_info=True)


class MarkdownConverter:

    def __init__(self, output_dir: Path) -> None:
        self.output_dir = output_dir

    def convert_md_to_docx(self, file_path: Path) -> Path:
        """Converts a Markdown file to a DOCX file using Pandoc.
        Args:file_path (Path): The path to the Markdown file.
        Returns:Path: The path to the converted DOCX file."""
        output_file = self.output_dir / f"{file_path.stem}.docx"
        pandoc_command = ["pandoc", str(file_path), "-o", str(output_file)]
        try:
            subprocess.run(pandoc_command, check=True)
            logger.info(f"Successfully converted {file_path} to {output_file}")
            return output_file
        except subprocess.CalledProcessError as e:
            logger.error(f"Error converting {file_path} to docx: {e}")
            return file_path


class DocumentFormatter:
    TOP_MARGIN = Cm(1)
    BOTTOM_MARGIN = Cm(1)
    LEFT_MARGIN = Cm(1)
    RIGHT_MARGIN = Cm(1)

    def __init__(self, doc) -> None:
        self.doc = doc

    def set_document_font(self, font_name='Open Sans', font_size=Pt(10)) -> None:
        try:
            for paragraph in self.doc.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = font_size
        except Exception as e:
            logger.error(f"Error setting document font: {e}", exc_info=True)

    def set_margins(self, top, bottom, left, right) -> None:
        for section in self.doc.sections:
            section.top_margin = top
            section.bottom_margin = bottom
            section.left_margin = left
            section.right_margin = right

    def modify_document_styles(self) -> None:
        try:
            styles = self.doc.styles
            default_style = styles['Normal']
            default_style.font.name = 'Open Sans'
            default_style.font.size = Pt(10)

            for style_name in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5']:
                if style_name in styles:
                    heading_style = styles[style_name]
                else:
                    heading_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                self._apply_style('Open Sans', heading_style, 16, 0, bold=True)

            heading1_style = styles['Heading 1'] if 'Heading 1' in styles else styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            self._apply_style('Open Sans', heading1_style, 16, 0)

            heading2_style = styles['Heading 2'] if 'Heading 2' in styles else styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            self._apply_style('Open Sans', heading2_style, 14, 5)

            block_text_style = styles['Block Text'] if 'Block Text' in styles else styles.add_style('Block Text', WD_STYLE_TYPE.PARAGRAPH)
            self._apply_style('Open Sans', block_text_style, 14, 5)
        except Exception as e:
            logger.error(f"Error modifying document styles: {e}", exc_info=True)

    def _apply_style(self, font_name, style, font_size, font_color_offset, bold=False) -> None:
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, font_color_offset, 255)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


class ImageResizer:

    def __init__(self, doc) -> None:
        self.doc = doc

    def autofit_images_to_window(self) -> None:
        default_width = Mm(210)  # A4 width in mm
        default_margin_left = Cm(1)
        default_margin_right = Cm(1)

        for section in self.doc.sections:
            page_width = section.page_width or default_width
            left_margin = section.left_margin or default_margin_left
            right_margin = section.right_margin or default_margin_right
            usable_width_emus = page_width - left_margin - right_margin

            for shape in self.doc.inline_shapes:
                if shape.type in (WD_INLINE_SHAPE.PICTURE, WD_INLINE_SHAPE.LINKED_PICTURE):
                    aspect_ratio = float(shape.height) / float(shape.width)
                    new_width = usable_width_emus
                    new_height = round(new_width * aspect_ratio)
                    shape.width = new_width
                    shape.height = new_height


class SectionManager:

    def __init__(self, doc) -> None:
        self.doc = doc

    def keep_sections_together(self) -> None:
        try:
            # Convert document to a single string for regex processing
            doc_text = '\n'.join([p.text for p in self.doc.paragraphs])
            # Find the TOC section using regular expressions
            start_marker = r"\s*# Table of Contents"  # ! here
            end_marker = r"\n\n---"
            start_match = re.search(start_marker, doc_text)
            end_match = re.search(end_marker, doc_text[start_match.end():], re.MULTILINE) if start_match else None
            # Identify the range of the TOC section
            if start_match and end_match:
                toc_section_start = start_match.start()
                toc_section_end = start_match.end() + end_match.end()
                for paragraph in self.doc.paragraphs:
                    paragraph_index = paragraph._element.getparent().getparent().index(paragraph._element)  # pylint: disable=protected-access
                    if toc_section_start <= paragraph_index <= toc_section_end:
                        continue  # Skip paragraphs in the TOC section
                    if paragraph.text.startswith("2. AWS") or paragraph.text.startswith("3. WPEngine"):
                        run = paragraph.add_run()
                        run.add_break(WD_BREAK.PAGE)
        except Exception as e:
            logger.error(f"Error keeping sections together: {e}", exc_info=True)

    def add_page_break_before_section(self, section_titles) -> None:
        try:
            toc_start, toc_end = self._find_toc_section()
            for paragraph in self.doc.paragraphs:
                paragraph_index = paragraph._element.getparent().index(paragraph._element)  # pylint: disable=protected-access
                if toc_start <= paragraph_index <= toc_end:
                    continue  # Skip paragraphs in the TOC section
                for title in section_titles:
                    if title in paragraph.text:
                        self._add_page_break_to_paragraph(paragraph.insert_paragraph_before())
                        break
        except Exception as e:
            logger.error(f"Error adding page break before section: {e}", exc_info=True, stack_info=True)

    def _find_toc_section(self) -> tuple[int, int] | tuple[Literal[-1], Literal[-1]]:
        doc_text = '\n'.join([p.text for p in self.doc.paragraphs])
        start_marker = r"\s*# Table of Contents"
        end_marker = r"\n\n---"
        start_match = re.search(start_marker, doc_text)
        end_match = re.search(end_marker, doc_text[start_match.end():], re.MULTILINE) if start_match else None
        if start_match and end_match:
            return start_match.start(), start_match.end() + end_match.end()
        return -1, -1

    def _add_page_break_to_paragraph(self, paragraph) -> None:
        try:
            run = paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)
        except Exception as e:
            logger.error(f"Error adding page break to paragraph: {e}", exc_info=True)


class StyleApplier:

    def __init__(self, doc, table_styler) -> None:
        self.doc = doc
        self.table_styler = table_styler

    def apply_custom_styles(self) -> None:
        try:
            for table in self.doc.tables:
                header_cells = table.rows[0].cells
                header_texts = [cell.text.strip() for cell in header_cells if cell.text.strip() != '']
                if CS.is_azure_table(header_texts):
                    CS.style_azure_table(self.table_styler, table)
                    logger.info("CS.style_azure_table(self.table_styler, table) done...")
                elif CS.is_wpengine_table(header_texts):
                    CS.style_wpengine_table(self.table_styler, table)
                    logger.info("CS.style_wpengine_table(self.table_styler, table) done...")
                elif CS.is_cisco_table(header_texts):
                    CS.style_cisco_table(self.table_styler, table)
                    logger.info("CS.style_cisco_table(self.table_styler, table) done...")
                elif CS.is_barracuda_table(header_texts):
                    CS.style_barracuda_table(self.table_styler, table)
                    logger.info("CS.style_barracuda_table(self.table_styler, table) done...")
                elif CS.is_websites_table(header_texts):
                    CS.style_websites_table(self.table_styler, table)
                    logger.info("CS.style_websites_table(self.table_styler, table) done...")
                elif CS.is_summary_table(header_texts):
                    CS.style_summary_table(self.table_styler, table)
                    logger.info("CS.style_summary_table(self.table_styler, table) done...")
        except Exception as e:
            logger.error(f"Error applying custom styles: {e}")


class TableStyler:
    TOP_MARGIN = Cm(1)
    BOTTOM_MARGIN = Cm(1)
    LEFT_MARGIN = Cm(1)
    RIGHT_MARGIN = Cm(1)

    def __init__(self, doc) -> None:
        self.doc = doc

    def autofit_tables_to_window(self) -> None:
        try:
            for table in self.doc.tables:
                default_width = Mm(210)  # A4 width in mm
                default_margin = Cm(1)
                section = self.doc.sections[0]
                page_width = section.page_width or default_width
                left_margin = section.left_margin or default_margin
                right_margin = section.right_margin or default_margin
                total_width = page_width - left_margin - right_margin
                for row in table.rows:
                    for cell in row.cells:
                        cell.width = int(total_width / len(row.cells))
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                self.apply_bold_to_headers(table)
        except Exception as e:
            logger.error(f"Error autofit_tables_to_window: {e}", exc_info=True)

    def apply_bold_to_headers(self, table):
        header_row = table.rows[0]
        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

    def style_table(self, table, header_fill, header_font_color, content_fill, content_font_color) -> None:
        try:
            for row_index, row in enumerate(table.rows):
                is_header_row = row_index == 0  # Identify header row
                fill_color = header_fill if is_header_row else content_fill
                font_color = header_font_color if is_header_row else content_font_color
                for cell in row.cells:
                    self.set_cell_borders(cell)
                    self.set_cell_background_color(cell, fill_color)
                    self.set_font_color(cell, font_color)
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Center-align vertically
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            run.font.name = 'Open Sans'
                            run.font.bold = is_header_row  # Apply bold only for header row
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Center-align vertically
                    self.set_cell_borders(cell)
                    self.set_cell_background_color(cell, header_fill if row == table.rows[0] else content_fill)
                    self.set_font_color(cell, header_font_color if row == table.rows[0] else content_font_color)
                    # Set paragraph alignment to center for horizontal centering
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Center-align vertically
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            run.font.name = 'Open Sans'
                            if is_header_row:
                                run.font.bold = True  # Bold for header row
                            else:
                                run.font.bold = False  # Regular for other rows
        except Exception as e:
            logger.error(f"Error style_table: {e}", exc_info=True)

    def set_cell_background_color(self, cell, color_str) -> None:
        try:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), color_str)
            cell._tc.get_or_add_tcPr().append(shading_elm)  # pylint: disable=protected-access
        except Exception as e:
            logger.error(f"Error set_cell_background_color: {e}", exc_info=True)

    def set_font_color(self, cell, font_color) -> None:
        try:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = font_color
        except Exception as e:
            logger.error(f"Error set_font_color: {e}", exc_info=True)

    def set_cell_borders(self, cell) -> None:
        try:
            tcPr = cell._tc.get_or_add_tcPr()  # pylint: disable=protected-access
            for border in ["top", "left", "bottom", "right"]:
                border_elm = OxmlElement(f'w:{border}')
                border_elm.set(qn('w:val'), 'single')
                border_elm.set(qn('w:sz'), '4')
                border_elm.set(qn('w:space'), '0')
                border_elm.set(qn('w:color'), 'auto')
                tcPr.append(border_elm)
        except Exception as e:
            logger.error(f"Error set_cell_borders: {e}", exc_info=True)

    def keep_table_together(self, table) -> None:
        try:
            tblPr = table._element.get_or_add_tblPr()  # pylint: disable=protected-access
            tblKeep = OxmlElement('w:tblpPr')
            tblKeep.set(qn('w:keepLines'), "1")
            tblPr.append(tblKeep)
            logger.info("keep_table_together completed.")
        except Exception as e:
            logger.error(f"Error keep_table_together: {e}", exc_info=True)

    # Helper function to create a qualified name (QName)
    def qname(self, tag) -> str:
        return f'{{{nsdecls("w").strip()}}}{tag}'

    def style_table_row(self, row, fill, font_color):
        for cell in row.cells:
            self.set_cell_background_color(cell, fill)  # Apply background color to each cell
            self.set_font_color(cell, font_color)  # Apply font color to each cell
            self.set_cell_borders(cell)  # Apply border styling to each cell

    def style_table_with_alternating_rows(self, table, header_fill, header_font_color, content_fill_1, content_fill_2, content_font_color) -> None:
        self.style_table_row(table.rows[0], header_fill, header_font_color)
        for i, row in enumerate(table.rows[1:]):
            fill = content_fill_1 if i % 2 == 0 else content_fill_2
            self.style_table_row(row, fill, content_font_color)


class CS():

    @staticmethod
    def style_table(doc_processor, table, header_fill, header_font_color, content_fill_1, content_fill_2, content_font_color):
        # Style the header row
        doc_processor.table_styler.style_table_row(table.rows[0], header_fill, header_font_color)
        # Style the content rows with alternating colors
        for i, row in enumerate(table.rows[1:]):
            fill = content_fill_1 if i % 2 == 0 else content_fill_2
            doc_processor.table_styler.style_table_row(row, fill, content_font_color)

    @staticmethod
    def style_azure_table(table_styler, table) -> None:
        try:
            azure_header_fill = '0078D7'  # Azure Blue
            azure_content_fill_1 = 'DEEBF7'  # Light Azure Blue
            azure_content_fill_2 = 'B3C6E7'  # Lighter Azure Blue
            azure_header_font_color = RGBColor(255, 255, 255)  # White
            azure_content_font_color = RGBColor(0, 0, 0)  # Black
            table_styler.style_table_with_alternating_rows(table, azure_header_fill, azure_header_font_color, azure_content_fill_1,
                                                           azure_content_fill_2, azure_content_font_color)
        except Exception as e:
            logger.error(f"Error style_azure_table: {e}", exc_info=True)

    @staticmethod
    def style_wpengine_table(table_styler, table) -> None:
        try:
            wpengine_header_fill = '8DB600'  # WP Engine Green
            wpengine_content_fill_1 = 'ECF0E7'  # Light WP Engine Green
            wpengine_content_fill_2 = 'D9EAD3'  # Lighter WP Engine Green
            wpengine_header_font_color = RGBColor(255, 255, 255)  # White
            wpengine_content_font_color = RGBColor(0, 0, 0)  # Black
            table_styler.style_table_with_alternating_rows(table, wpengine_header_fill, wpengine_header_font_color, wpengine_content_fill_1,
                                                           wpengine_content_fill_2, wpengine_content_font_color)
        except Exception as e:
            logger.error(f"Error occured in styles {e}")

    @staticmethod
    def style_cisco_table(table_styler, table) -> None:
        try:
            # cisco_header_fill = '2E8B57'  # Sea Green
            # cisco_content_fill_1 = '66CDAA'  # Medium Aquamarine
            # cisco_content_fill_2 = '98FB98'  # Pale Green
            cisco_header_fill = '86A697'  # Muted Sea Green
            cisco_content_fill_1 = 'ACC6B5'  # Soft Green
            cisco_content_fill_2 = 'CFDED6'  # Pale Green
            cisco_header_font_color = RGBColor(0, 0, 0)  # Black
            cisco_content_font_color = RGBColor(0, 0, 0)  # Black
            table_styler.style_table_with_alternating_rows(table, cisco_header_fill, cisco_header_font_color, cisco_content_fill_1,
                                                           cisco_content_fill_2, cisco_content_font_color)
        except Exception as e:
            logger.error(f"Error occured in styles {e}")

    @staticmethod
    def style_barracuda_table(table_styler, table) -> None:
        try:
            barracuda_header_fill = '006888'  # Barracuda Blue
            barracuda_content_fill_1 = 'E1EFF6'  # Light Barracuda Blue
            barracuda_content_fill_2 = 'D9E8F2'  # Lighter Barracuda Blue
            barracuda_header_font_color = RGBColor(255, 255, 255)  # White
            barracuda_content_font_color = RGBColor(0, 0, 0)  # Black
            table_styler.style_table_with_alternating_rows(table, barracuda_header_fill, barracuda_header_font_color, barracuda_content_fill_1,
                                                           barracuda_content_fill_2, barracuda_content_font_color)
        except Exception as e:
            logger.error(f"Error occured in styles {e}")

    @staticmethod
    def style_websites_table(table_styler, table) -> None:
        try:
            # websites_header_fill = '4682B4'  # Steel Blue
            # websites_content_fill_1 = 'B0C4DE'  # Light Steel Blue
            # websites_content_fill_2 = '87CEFA'  # Light Sky Blue
            websites_header_fill = '7A9DAB'  # Dusty Blue
            websites_content_fill_1 = 'A8C0CF'  # Soft Blue
            websites_content_fill_2 = 'D0E1EC'  # Pale Blue
            websites_header_font_color = RGBColor(0, 0, 0)  # Black
            websites_content_font_color = RGBColor(0, 0, 0)  # Black
            table_styler.style_table_with_alternating_rows(table, websites_header_fill, websites_header_font_color, websites_content_fill_1,
                                                           websites_content_fill_2, websites_content_font_color)
        except Exception as e:
            logger.error(f"Error occured in styles {e}")

    @staticmethod
    def style_summary_table(table_styler, table) -> None:
        try:
            summary_header_fill = 'FFBF00'  # Amber
            summary_content_fill_1 = 'FFD700'  # Gold
            summary_content_fill_2 = 'FFECB3'  # Pale Gold

            # summary_header_fill = '8B4513'  # Saddle Brown
            # summary_content_fill_1 = 'CD853F'  # Peru
            # summary_content_fill_2 = 'DEB887'  # Burlywood
            summary_header_font_color = RGBColor(0, 0, 0)  # Black
            summary_content_font_color = RGBColor(0, 0, 0)  # Black
            table_styler.style_table_with_alternating_rows(table, summary_header_fill, summary_header_font_color, summary_content_fill_1,
                                                           summary_content_fill_2, summary_content_font_color)
        except Exception as e:
            logger.error(f"Error occured in styles {e}")

    @staticmethod
    def is_summary_table(header_texts) -> bool:
        try:
            summary_header = ["Business", "Coding", "Item", "Notes", "Status"]
            # Allow partial matching for headers, ignore case and whitespace differences
            return all(any(h.strip().lower() == expected_header.lower() for h in header_texts) for expected_header in summary_header)
        except Exception as e:
            logger.error(f"Error is_summary_table: {e}", exc_info=True)
            return False

    @staticmethod
    def is_barracuda_table(header_texts) -> bool:
        try:
            barracuda_texts = [["Corporate", "Email Blocked", "BRBL", "SPAM", "BRTS", "Virus", "ATP"],
                               ["Payments", "Email Blocked", "BRBL", "SPAM", "BRTS", "Virus", "ATP"],
                               ["Prepaid", "Email Blocked", "BRBL", "SPAM", "BRTS", "Virus", "ATP"],
                               ["SmartCentral", "Email Blocked", "BRBL", "SPAM", "BRTS", "Virus", "ATP"],
                               ["Summary", "Email Blocked", "BRBL", "SPAM", "BRTS", "Virus", "ATP", "Blocked Email%", "Blocked ATP%"]]
            return any(header_texts == barracuda_header for barracuda_header in barracuda_texts)
        except Exception as e:
            logger.error(f"Error is_barracuda_table: {e}", exc_info=True)
            return False

    @staticmethod
    def is_websites_table(header_texts) -> bool:
        try:
            websites_texts = [["Corporate", "Avg daily traffic", "WPScan Vulns", "Site WAF", "Plugins", "Themes", "WP ver", "PHP ver"],
                              ["Payments", "Avg daily traffic", "WPScan Vulns", "Site WAF", "Plugins", "Themes", "WP ver", "PHP ver"],
                              ["Prepaid", "Avg daily traffic", "WPScan Vulns", "Site WAF", "Plugins", "Themes", "WP ver", "PHP ver"],
                              ["SmartCentral", "Avg daily traffic", "WPScan Vulns", "Site WAF", "Plugins", "Themes", "WP ver", "PHP ver"]]
            return any(header_texts == websites_header for websites_header in websites_texts)
        except Exception as e:
            logger.error(f"Error is_websites_table: {e}", exc_info=True)
            return False

    @staticmethod
    def is_wpengine_table(header_texts) -> bool:
        wpengine_texts = ["Plugins updated", "Domains secured", "Platform enhancements", "Attacks blocked"]
        return header_texts == wpengine_texts

    @staticmethod
    def is_cisco_table(header_texts) -> bool:
        try:
            cisco_headers = [[
                "Total Data Transferred", "Total Data - DOWNLOADED", "Total Data - UPLOADED", "Total Unique Clients", "Average of clients per day",  # pylint: disable=line-too-long
                "Average usage per client"
            ], ["Top clients by usage", "Usage", "Usage", "Top Blocked Sites by URL", "Category", "Sites"]]  # pylint: disable=line-too-long
            # logger.info(f"cisco_headers: {cisco_headers}")
            return header_texts in cisco_headers
        except Exception as e:
            logger.error(f"Error is_cisco_table: {e}", exc_info=True)
            return False

    @staticmethod
    def is_azure_table(header_texts) -> bool:
        azure_headers = [["Failing Controls - UGC", "Failing Controls - ZenPay"], ["Control States:", "UGC", "ZenPay"],
                         ["Resource States:", "UGC", "ZenPay"]]
        azure_other_header = len(header_texts) == 6 and header_texts[3] == ''
        return header_texts in azure_headers or azure_other_header


install()


def configure_logging() -> None:
    logging.basicConfig(level=logging.INFO,
                        format="%(message)s",
                        handlers=[
                            RichHandler(level=logging.INFO,
                                        show_time=True,
                                        show_path=True,
                                        show_level=True,
                                        rich_tracebacks=True,
                                        tracebacks_extra_lines=0,
                                        tracebacks_show_locals=False)
                        ])


logger = logging.getLogger(__name__)


def main() -> None:
    configure_logging()
    dir_init = DirectoryInitializer()
    markdown_converter = MarkdownConverter(dir_init.output_dir)
    files_converted = False

    for file_path in dir_init.input_dir.iterdir():
        if file_path.suffix == ".md":
            logger.info(f'Processing Markdown file: {file_path}')
            if docx_file_path := markdown_converter.convert_md_to_docx(file_path):
                doc_processor = DocxProcessor(str(dir_init.input_dir), str(dir_init.output_dir), str(dir_init.reference_dir),
                                              str(dir_init.reference_doc))
                doc_processor.post_process_docx(docx_file_path)
                files_converted = True

    if not files_converted:
        logger.info("No Markdown files found to convert.")


if __name__ == "__main__":
    main()
